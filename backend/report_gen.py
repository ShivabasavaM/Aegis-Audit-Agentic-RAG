from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_docx_report(audit_results, metadata):
    doc = Document()
    
    title = doc.add_heading('Aegis-Audit: Legal Verification Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Executive Summary', level=1)
    p = doc.add_paragraph()
    p.add_run("Reference Law: ").bold = True
    p.add_run(f"{metadata.get('law_name', 'N/A')}\n")
    p.add_run("Internal Policy: ").bold = True
    p.add_run(f"{metadata.get('policy_name', 'N/A')}")

    doc.add_heading('2. Detailed Findings & Remediation', level=1)

    for item in audit_results:
        pillar_name = item.get('pillar', 'General Pillar')
        doc.add_heading(pillar_name, level=2)
        
        rating = item.get('rating', 'Low')
        r_para = doc.add_paragraph()
        run = r_para.add_run(f"RISK LEVEL: {rating}")
        run.bold = True
        if rating in ["Critical", "High"]:
            run.font.color.rgb = RGBColor(255, 0, 0)  
        elif rating == "Medium":
            run.font.color.rgb = RGBColor(255, 165, 0) 
            
        finding_para = doc.add_paragraph()
        finding_para.add_run("Finding: ").bold = True
        finding_para.add_run(item.get('finding', 'No specific finding recorded.'))
        
        meta_para = doc.add_paragraph()
        source_run = meta_para.add_run(f"Source: {item.get('citation', 'Not cited')}")
        source_run.italic = True
        source_run.font.size = Pt(9)
        
        conf_run = meta_para.add_run(f" | AI Confidence: {item.get('confidence', 'N/A')}")
        conf_run.font.size = Pt(9)

        doc.add_heading("Remediation Plan:", level=3)
        doc.add_paragraph(item.get('remediation', 'No remediation required.'))
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer